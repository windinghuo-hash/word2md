from typing import List, Optional
import os
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn
from utils.ast_nodes import Node, HeadingNode, ParagraphNode, ListNode, CodeBlockNode, TableNode, ImageNode

# 字体大小（pt）到标题层级的映射，用于处理无标准样式的文档
_SIZE_TO_HEADING = {
    26: 1,
    24: 1,
    22: 1,
    20: 2,
    18: 2,
    16: 3,
    15: 4,
    14: 5,
    13: 6,
    12: 6,
}

# Word 图片关系类型
_IMAGE_RTYPE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'

class DocxParser:
    def parse(self, file_path: str, extract_images: bool = False, images_dir: str = None) -> List[Node]:
        self._doc_path = file_path
        self._extract_images = extract_images
        self._images_dir = images_dir
        self._image_counter = 0

        doc = Document(file_path)
        self._doc = doc
        nodes = []

        # 按文档真实顺序遍历段落和表格
        for element in doc.element.body:
            tag = element.tag.split('}')[-1]
            if tag == 'p':
                para = Paragraph(element, doc)
                node = self._parse_paragraph(para)
                if node:
                    if isinstance(node, list):
                        nodes.extend(node)
                    else:
                        nodes.append(node)
            elif tag == 'tbl':
                table = Table(element, doc)
                node = self._parse_table(table)
                if node:
                    nodes.append(node)

        return nodes

    def _parse_paragraph(self, para) -> Optional[Node]:
        # 先检查段落内是否包含图片
        images = self._extract_images_from_para(para)
        if images:
            # 段落内有图片，返回图片节点列表（可能还有文字，拆开处理）
            nodes = []
            nodes.extend(images)
            # 如果段落还有文字，也保留
            text = para.text.strip()
            if text:
                nodes.append(ParagraphNode(runs=[{"text": text, "bold": False, "italic": False}]))
            return nodes  # 返回列表，由 parse() 展开

        # 跳过空段落
        if not para.text.strip():
            return None

        style_name = para.style.name if para.style else ""

        # 标准 Heading 样式
        if style_name.startswith('Heading'):
            try:
                level = int(style_name.split()[-1])
                return HeadingNode(level=level, content=para.text)
            except (ValueError, IndexError):
                pass

        # 标准列表样式
        if 'List' in style_name:
            return ListNode(items=[para.text], ordered='Number' in style_name)

        # 代码块（等宽字体）
        if self._is_code_block(para):
            return CodeBlockNode(code=para.text, language="")

        # 无标准样式时，通过字体大小+粗体推断标题
        heading_level = self._infer_heading_level(para)
        if heading_level:
            return HeadingNode(level=heading_level, content=para.text)

        # 普通段落
        runs = []
        for run in para.runs:
            if not run.text:
                continue
            runs.append({
                "text": run.text,
                "bold": run.bold or False,
                "italic": run.italic or False
            })
        if runs:
            return ParagraphNode(runs=runs)
        return None

    def _extract_images_from_para(self, para) -> List[ImageNode]:
        """从段落中提取所有图片，返回 ImageNode 列表"""
        image_nodes = []

        # 查找段落中所有 <a:blip> 元素（图片引用）
        drawing_elements = para._element.findall('.//' + qn('a:blip'))
        if not drawing_elements:
            return []

        for blip in drawing_elements:
            embed_id = blip.get(qn('r:embed'))
            if not embed_id:
                continue

            try:
                part = self._doc.part.related_parts.get(
                    self._doc.part.rels[embed_id].target_ref
                )
                if part is None:
                    # 通过 rId 查找
                    rel = self._doc.part.rels.get(embed_id)
                    if rel is None:
                        continue
                    part = rel.target_part
            except (KeyError, AttributeError):
                continue

            self._image_counter += 1
            ext = os.path.splitext(part.partname)[1] or '.png'
            filename = f"image_{self._image_counter:03d}{ext}"

            if self._extract_images and self._images_dir:
                # 保存图片到磁盘
                os.makedirs(self._images_dir, exist_ok=True)
                img_path = os.path.join(self._images_dir, filename)
                with open(img_path, 'wb') as f:
                    f.write(part.blob)
                # 使用相对路径（相对于输出 md 文件）
                rel_path = os.path.join(os.path.basename(self._images_dir), filename).replace('\\', '/')
                image_nodes.append(ImageNode(path=rel_path, alt=f"图片{self._image_counter}"))
            else:
                # 不提取，仅标记占位符
                image_nodes.append(ImageNode(path="", alt=f"图片{self._image_counter}（未提取）"))

        return image_nodes

    def _infer_heading_level(self, para) -> Optional[int]:
        """通过字体大小和粗体属性推断标题层级"""
        if not para.runs:
            return None
        run = para.runs[0]
        is_bold = run.bold
        size = run.font.size
        if not size or not is_bold:
            return None
        size_pt = round(size / 12700)  # EMU 转 pt
        return _SIZE_TO_HEADING.get(size_pt)

    def _is_code_block(self, para) -> bool:
        if not para.runs:
            return False
        font_name = para.runs[0].font.name
        return font_name in ['Courier New', 'Consolas', 'Monaco']

    def _parse_table(self, table: Table) -> Optional[Node]:
        """解析表格为 TableNode 或 CodeBlockNode"""
        if not table.rows:
            return None

        # 检测单列表格（通常是代码块/流程图）
        col_count = len(table.columns)
        if col_count == 1:
            # 单列表格作为代码块处理
            lines = []
            for row in table.rows:
                if row.cells:
                    cell_text = '\n'.join(p.text for p in row.cells[0].paragraphs if p.text.strip())
                    if cell_text:
                        lines.append(cell_text)

            if not lines:
                return None

            code = '\n'.join(lines)

            # 尝试提取语言标识（第一行可能是语言名）
            language = ""
            if lines and lines[0].strip() in ['sql', 'python', 'javascript', 'java', 'Plain Text', 'bash', 'json', 'yaml']:
                language = lines[0].strip()
                if language == 'Plain Text':
                    language = ""
                code = '\n'.join(lines[1:])  # 去掉语言标识行

            return CodeBlockNode(code=code, language=language)

        # 多列表格正常处理
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                # 提取单元格文本，保留段落间的换行
                cell_text = '\n'.join(p.text for p in cell.paragraphs if p.text.strip())
                cells.append(cell_text)
            rows.append(cells)

        # 判断是否有表头（第一行是否加粗）
        has_header = False
        if rows and table.rows[0].cells:
            first_cell = table.rows[0].cells[0]
            if first_cell.paragraphs:
                first_para = first_cell.paragraphs[0]
                if first_para.runs:
                    has_header = first_para.runs[0].bold or False

        return TableNode(rows=rows, has_header=has_header)
