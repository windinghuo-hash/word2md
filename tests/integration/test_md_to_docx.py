import pytest
import os
from docx import Document
from converters.md_to_docx import MarkdownToDocxConverter

def test_convert_markdown_to_docx():
    converter = MarkdownToDocxConverter()
    output_path = 'tests/fixtures/test_output.docx'
    converter.convert('tests/fixtures/test_basic.md', output_path)
    assert os.path.exists(output_path)
    doc = Document(output_path)
    assert len(doc.paragraphs) > 0
    headings = [p for p in doc.paragraphs if 'Heading' in p.style.name]
    assert len(headings) >= 2
    os.remove(output_path)
