import pytest
import os
from docx import Document
from renderers.docx_renderer import DocxRenderer
from utils.ast_nodes import HeadingNode, ParagraphNode, ListNode, CodeBlockNode

def test_render_heading():
    renderer = DocxRenderer()
    nodes = [HeadingNode(level=1, content="Test Title")]
    output_path = 'tests/fixtures/test_render.docx'
    renderer.render(nodes, output_path)
    doc = Document(output_path)
    assert len(doc.paragraphs) > 0
    assert doc.paragraphs[0].text == "Test Title"
    assert 'Heading 1' in doc.paragraphs[0].style.name
    os.remove(output_path)

def test_render_paragraph_with_bold():
    renderer = DocxRenderer()
    runs = [
        {"text": "Normal ", "bold": False, "italic": False},
        {"text": "bold", "bold": True, "italic": False}
    ]
    nodes = [ParagraphNode(runs=runs)]
    output_path = 'tests/fixtures/test_render.docx'
    renderer.render(nodes, output_path)
    doc = Document(output_path)
    para = doc.paragraphs[0]
    assert len(para.runs) == 2
    assert para.runs[1].bold is True
    os.remove(output_path)
