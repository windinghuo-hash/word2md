from typing import List
from docx import Document
from docx.shared import Pt
from utils.ast_nodes import Node

class DocxRenderer:
    def render(self, nodes: List[Node], output_path: str):
        doc = Document()
        for node in nodes:
            self._render_node(doc, node)
        doc.save(output_path)

    def _render_node(self, doc: Document, node: Node):
        if node.type == "heading":
            self._render_heading(doc, node)
        elif node.type == "paragraph":
            self._render_paragraph(doc, node)
        elif node.type == "list":
            self._render_list(doc, node)
        elif node.type == "code_block":
            self._render_code_block(doc, node)

    def _render_heading(self, doc, node):
        doc.add_heading(node.content, level=node.attrs["level"])

    def _render_paragraph(self, doc, node):
        para = doc.add_paragraph()
        for run_data in node.content:
            run = para.add_run(run_data["text"])
            run.bold = run_data.get("bold", False)
            run.italic = run_data.get("italic", False)

    def _render_list(self, doc, node):
        ordered = node.attrs["ordered"]
        style = 'List Number' if ordered else 'List Bullet'
        for item in node.content:
            doc.add_paragraph(item, style=style)

    def _render_code_block(self, doc, node):
        para = doc.add_paragraph(node.content)
        if para.runs:
            para.runs[0].font.name = 'Courier New'
            para.runs[0].font.size = Pt(10)
